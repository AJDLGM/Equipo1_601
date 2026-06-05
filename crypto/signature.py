import os
import re
import json
import zipfile
import hashlib
import base64
import io
import shutil
from datetime import datetime
try:
    from zoneinfo import ZoneInfo as _ZoneInfo
except ImportError:
    from backports.zoneinfo import ZoneInfo as _ZoneInfo
_CDMX = _ZoneInfo("America/Mexico_City")
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import padding
from config.paths import get_user_dir

_W_NS           = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_DOCX_SIG_MARK  = "%%FIRMA-DIGITAL"      # párrafo invisible que delimita inicio de firma
_DOCX_JSON_MARK = "%%FIRMA-DIGITAL-JSON%%:"  # párrafo con los datos de verificación
_PDF_META_KEY   = "/FirmaDigital"         # clave en metadatos PDF
_PDF_JSON_MARK  = "%%FIRMA-DIGITAL-JSON%%:"


# ── Helpers criptográficos ────────────────────────────────────────────────────

def _load_private_key(username):
    dirs = get_user_dir(username)
    path = os.path.join(dirs["keys"], f"{username}_private.pem")
    with open(path, "rb") as f:
        return serialization.load_pem_private_key(f.read(), password=None)


def _load_certificate(username):
    dirs = get_user_dir(username)
    path = os.path.join(dirs["certs"], f"{username}_cert.json")
    with open(path, "r", encoding="utf-8") as f:
        cert = json.load(f)

    # Certs en caché antiguo no tienen public_key — lo tomamos del .pem local
    if not cert.get("public_key"):
        pub_path = os.path.join(dirs["keys"], f"{username}_public.pem")
        if os.path.exists(pub_path):
            with open(pub_path, "r", encoding="utf-8") as f:
                cert["public_key"] = f.read()

    return cert


def _rsa_sign(private_key, data: bytes) -> bytes:
    return private_key.sign(
        data,
        padding.PSS(
            mgf=padding.MGF1(hashes.SHA256()),
            salt_length=padding.PSS.MAX_LENGTH,
        ),
        hashes.SHA256(),
    )


def _rsa_verify(public_key, signature: bytes, data: bytes):
    public_key.verify(
        signature,
        data,
        padding.PSS(
            mgf=padding.MGF1(hashes.SHA256()),
            salt_length=padding.PSS.MAX_LENGTH,
        ),
        hashes.SHA256(),
    )


def _public_key_from_cert(certificate: dict):
    pem = certificate.get("public_key", "")
    if not pem:
        raise ValueError("El certificado no contiene clave publica.")
    return serialization.load_pem_public_key(pem.encode())


def _build_sig_meta(username, file_path, file_hash, raw_sig, certificate):
    return {
        "signer":            username,
        "original_filename": os.path.basename(file_path),
        "algorithm":         "RSA-PSS-SHA256",
        "hash_sha256":       file_hash,
        "signature":         base64.b64encode(raw_sig).decode(),
        "certificate":       certificate,
        "signed_at":         datetime.now(_CDMX).strftime("%Y-%m-%dT%H:%M:%S"),
    }


def _info_from_meta(sig_meta):
    cert = sig_meta.get("certificate", {})
    return {
        "signer":            sig_meta.get("signer", "—"),
        "original_filename": sig_meta.get("original_filename", "—"),
        "signed_at":         sig_meta.get("signed_at", ""),
        "hash_sha256":       sig_meta.get("hash_sha256", ""),
        "cert_status":       cert.get("status", "unknown"),
        "cert_issued":       cert.get("issued_at", ""),
        "cert_expires":      cert.get("expires_at", ""),
    }


def _to_b64_json(obj) -> str:
    return base64.b64encode(
        json.dumps(obj, ensure_ascii=False).encode("utf-8")
    ).decode()


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _get_para_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.iter(f"{_W_NS}t"))


def _docx_content_hash(doc_path) -> str:
    """
    Hash SHA-256 determinista del cuerpo del documento DOCX.
    Itera todos los hijos directos de <w:body> (párrafos, tablas, etc.)
    excluyendo <w:sectPr> y deteniéndose al encontrar el marcador de firma.
    Al ser XML serializado, captura texto Y formato.
    """
    from docx import Document
    from lxml import etree

    doc = Document(doc_path)
    h   = hashlib.sha256()
    for child in doc.element.body:
        if child.tag == f"{_W_NS}sectPr":          # propiedades de sección, se excluyen
            continue
        if (child.tag == f"{_W_NS}p"
                and _get_para_text(child).startswith(_DOCX_SIG_MARK)):
            break                                   # inicio de la sección de firma
        h.update(etree.tostring(child, encoding="unicode").encode("utf-8"))
    return h.hexdigest()


def _append_docx_sig_block(doc, sig_meta, certificate):
    """Añade un bloque de firma a un objeto Document ya abierto."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    signed_at = sig_meta["signed_at"][:19].replace("T", " ")
    username  = sig_meta.get("signer", "")
    order     = sig_meta.get("order")
    is_route  = sig_meta.get("route_signature", False)

    p   = doc.add_paragraph()
    run = p.add_run(_DOCX_SIG_MARK)
    run.font.size      = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    sep = doc.add_paragraph("_" * 72)
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after  = Pt(6)

    title = f"FIRMA DIGITAL #{order}" if (is_route and order) else "FIRMA DIGITAL"
    p   = doc.add_paragraph()
    run = p.add_run(title)
    run.bold      = True
    run.font.size = Pt(14)
    p.alignment   = WD_ALIGN_PARAGRAPH.CENTER

    fields = [
        ("Firmante",           username),
        ("Algoritmo",          "RSA-PSS-SHA256"),
        ("Fecha de firma",     signed_at),
        ("Estado cert.",       certificate.get("status", "active").upper()),
        ("Cert. válido hasta", certificate.get("expires_at", "")[:19]),
        ("Hash SHA-256",       sig_meta.get("hash_sha256", "")[:48] + "..."),
    ]
    for label, value in fields:
        p     = doc.add_paragraph()
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        p.add_run(str(value))

    p   = doc.add_paragraph()
    run = p.add_run(f"{_DOCX_JSON_MARK}{_to_b64_json(sig_meta)}")
    run.font.size      = Pt(5)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _sign_docx(username, file_path, private_key, certificate):
    from docx import Document

    file_hash = _docx_content_hash(file_path)
    raw_sig   = _rsa_sign(private_key, file_hash.encode())
    sig_meta  = _build_sig_meta(username, file_path, file_hash, raw_sig, certificate)

    base_, _ = os.path.splitext(file_path)
    signed_path = base_ + "_firmado.docx"
    shutil.copy2(file_path, signed_path)

    doc = Document(signed_path)
    _append_docx_sig_block(doc, sig_meta, certificate)
    doc.save(signed_path)
    return signed_path


def _verify_docx(file_path):
    from docx import Document

    doc = Document(file_path)

    # Recoger TODAS las firmas embebidas (orden de aparición = orden de firma)
    all_sigs = []
    for para in doc.paragraphs:
        if para.text.startswith(_DOCX_JSON_MARK):
            try:
                meta = json.loads(
                    base64.b64decode(para.text[len(_DOCX_JSON_MARK):]).decode("utf-8")
                )
                all_sigs.append(meta)
            except Exception:
                pass

    if not all_sigs:
        return False, "El documento no contiene firma digital embebida."

    all_sigs.sort(key=lambda s: s.get("order", 1))
    original_hash = _docx_content_hash(file_path)

    # ── Firma simple (formato original) ──────────────────────
    if len(all_sigs) == 1 and not all_sigs[0].get("route_signature"):
        sig_meta    = all_sigs[0]
        stored_hash = sig_meta.get("hash_sha256", "")
        if original_hash != stored_hash:
            return False, (
                "El contenido del documento fue modificado después de ser firmado.\n"
                "Los hashes no coinciden."
            )
        try:
            pub_key = _public_key_from_cert(sig_meta.get("certificate", {}))
            raw_sig = base64.b64decode(sig_meta["signature"])
            _rsa_verify(pub_key, raw_sig, stored_hash.encode())
        except ValueError as e:
            return False, str(e)
        except Exception:
            return False, "La firma digital no es válida."
        return True, _info_from_meta(sig_meta)

    # ── Ruta multi-firma ─────────────────────────────────────
    for i, sig in enumerate(all_sigs):
        # hash_chain esperado para este paso
        if i == 0:
            expected_hash = original_hash          # primer firmante: igual al hash del contenido
        else:
            chain_input = original_hash.encode()
            for prev in all_sigs[:i]:
                chain_input += base64.b64decode(prev["signature"])
            expected_hash = hashlib.sha256(chain_input).hexdigest()

        stored_hash = sig.get("hash_chain", sig.get("hash_sha256", ""))
        if stored_hash != expected_hash:
            return False, (
                f"La firma #{sig.get('order', i+1)} ({sig.get('signer','?')}) no es válida.\n"
                "El hash de cadena no coincide: orden alterado o firma faltante."
            )
        try:
            pub_key = _public_key_from_cert(sig.get("certificate", {}))
            raw_sig = base64.b64decode(sig["signature"])
            _rsa_verify(pub_key, raw_sig, expected_hash.encode())
        except ValueError as e:
            return False, str(e)
        except Exception:
            return False, (
                f"La firma #{sig.get('order', i+1)} ({sig.get('signer','?')}) "
                "no es válida criptográficamente."
            )

    last        = all_sigs[-1]
    all_signers = " → ".join(
        f"{s.get('signer','?')} (#{s.get('order', i+1)})"
        for i, s in enumerate(all_sigs)
    )
    return True, {
        "signer":            last.get("signer", ""),
        "original_filename": last.get("original_filename", os.path.basename(file_path)),
        "signed_at":         last.get("signed_at", ""),
        "hash_sha256":       original_hash,
        "cert_status":       last.get("certificate", {}).get("status", "unknown"),
        "route":             [],
        "all_signers":       all_signers,
        "total_signatures":  len(all_sigs),
    }


# ── PDF ───────────────────────────────────────────────────────────────────────

def _pdf_content_hash(pdf_path, skip_last_n: int = 0) -> str:
    """
    Hash SHA-256 del texto extraído de todas las páginas,
    omitiendo las últimas skip_last_n (la página de firma).
    """
    from pypdf import PdfReader
    reader = PdfReader(pdf_path)
    pages  = list(reader.pages)
    if skip_last_n:
        pages = pages[:-skip_last_n]
    h = hashlib.sha256()
    for page in pages:
        h.update((page.extract_text() or "").encode("utf-8"))
    return h.hexdigest()


def _create_pdf_sig_page(sig_meta: dict, certificate: dict) -> bytes:
    """Genera una página PDF con la sección de firma usando reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    buf    = io.BytesIO()
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "SigTitle", parent=styles["Heading1"],
        fontSize=16, alignment=TA_CENTER, spaceAfter=10,
    )
    small_style = ParagraphStyle(
        "SigSmall", parent=styles["Normal"],
        fontSize=5.5, textColor=colors.HexColor("#aaaaaa"),
        wordWrap="CJK", leading=7,
    )
    field_style = styles["Normal"]

    signed_at = sig_meta["signed_at"][:19].replace("T", " ")
    signer    = sig_meta.get("signer", "")
    file_hash = sig_meta.get("hash_sha256", "")

    story = [
        HRFlowable(width="100%", thickness=1.2, color=colors.HexColor("#aaaaaa")),
        Spacer(1, 0.18 * inch),
        Paragraph("FIRMA DIGITAL", title_style),
        Paragraph(f"<b>Firmante:</b> {signer}", field_style),
        Paragraph(f"<b>Algoritmo:</b> RSA-PSS-SHA256", field_style),
        Paragraph(f"<b>Fecha de firma:</b> {signed_at}", field_style),
        Paragraph(
            f"<b>Estado certificado:</b> {certificate.get('status','active').upper()}",
            field_style,
        ),
        Paragraph(
            f"<b>Cert. válido hasta:</b> {certificate.get('expires_at','')[:19]}",
            field_style,
        ),
        Paragraph(f"<b>Hash SHA-256:</b> {file_hash[:48]}...", field_style),
        Spacer(1, 0.18 * inch),
        HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd")),
        Spacer(1, 0.08 * inch),
        # Datos de verificación (pequeño, gris, visible)
        Paragraph(f"{_PDF_JSON_MARK}{_to_b64_json(sig_meta)}", small_style),
    ]

    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=inch, rightMargin=inch,
    )
    doc.build(story)
    return buf.getvalue()


def _sign_pdf(username, file_path, private_key, certificate):
    from pypdf import PdfReader, PdfWriter

    # Paso 1: hash del contenido original (texto de todas las páginas)
    file_hash = _pdf_content_hash(file_path)

    # Paso 2: firmar el hash
    raw_sig  = _rsa_sign(private_key, file_hash.encode())
    sig_meta = _build_sig_meta(username, file_path, file_hash, raw_sig, certificate)

    # Paso 3: crear la página de firma
    sig_page_bytes = _create_pdf_sig_page(sig_meta, certificate)

    # Paso 4: unir el PDF original + página de firma
    writer = PdfWriter()
    reader = PdfReader(file_path)
    for page in reader.pages:
        writer.add_page(page)

    sig_reader = PdfReader(io.BytesIO(sig_page_bytes))
    for page in sig_reader.pages:
        writer.add_page(page)

    # También en metadatos para lectura confiable
    writer.add_metadata({_PDF_META_KEY: _to_b64_json(sig_meta)})

    base_, _ = os.path.splitext(file_path)
    signed_path = base_ + "_firmado.pdf"
    with open(signed_path, "wb") as f:
        writer.write(f)

    return signed_path


def _verify_pdf(file_path):
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    if not reader.pages:
        return False, "El PDF no tiene paginas."

    # Intentar metadatos primero (más confiable)
    sig_meta = None
    meta = reader.metadata or {}
    if _PDF_META_KEY in meta:
        try:
            sig_meta = json.loads(base64.b64decode(meta[_PDF_META_KEY]).decode("utf-8"))
        except Exception:
            pass

    # Fallback: buscar el marcador en el texto de la última página
    if sig_meta is None:
        last_text = reader.pages[-1].extract_text() or ""
        idx = last_text.find(_PDF_JSON_MARK)
        if idx == -1:
            return False, "El PDF no contiene firma digital embebida."
        b64_data = "".join(last_text[idx + len(_PDF_JSON_MARK):].split())
        try:
            sig_meta = json.loads(base64.b64decode(b64_data).decode("utf-8"))
        except Exception:
            return False, "No se pudo decodificar la firma del PDF."

    # Recalcular hash del contenido original (todas las páginas excepto la última = firma)
    computed_hash = _pdf_content_hash(file_path, skip_last_n=1)
    stored_hash   = sig_meta.get("hash_sha256", "")
    if computed_hash != stored_hash:
        return False, (
            "El contenido del PDF fue modificado despues de ser firmado.\n"
            "Los hashes no coinciden."
        )

    # Verificar firma RSA con la clave pública del certificado embebido
    try:
        certificate = sig_meta.get("certificate", {})
        public_key  = _public_key_from_cert(certificate)
        raw_sig     = base64.b64decode(sig_meta["signature"])
        _rsa_verify(public_key, raw_sig, stored_hash.encode())
    except ValueError as e:
        return False, str(e)
    except Exception:
        return False, "La firma digital no es valida."

    return True, _info_from_meta(sig_meta)


# ── Otros formatos: contenedor .signed ───────────────────────────────────────

def _sign_container(username, file_path, private_key, certificate):
    with open(file_path, "rb") as f:
        file_bytes = f.read()

    file_hash = hashlib.sha256(file_bytes).hexdigest()
    raw_sig   = _rsa_sign(private_key, file_hash.encode())
    sig_meta  = _build_sig_meta(username, file_path, file_hash, raw_sig, certificate)

    signed_path = file_path + ".signed"
    with zipfile.ZipFile(signed_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("content", file_bytes)
        zf.writestr("signature.json",
                    json.dumps(sig_meta, indent=2, ensure_ascii=False))
    return signed_path


def _verify_container(file_path):
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            file_bytes = zf.read("content")
            sig_meta   = json.loads(zf.read("signature.json"))
    except Exception:
        return False, "El archivo no es un contenedor firmado valido."

    computed_hash = hashlib.sha256(file_bytes).hexdigest()
    stored_hash   = sig_meta.get("hash_sha256", "")
    if computed_hash != stored_hash:
        return False, "El archivo fue modificado despues de ser firmado."

    try:
        certificate = sig_meta.get("certificate", {})
        public_key  = _public_key_from_cert(certificate)
        raw_sig     = base64.b64decode(sig_meta["signature"])
        _rsa_verify(public_key, raw_sig, stored_hash.encode())
    except ValueError as e:
        return False, str(e)
    except Exception:
        return False, "La firma digital no es valida."

    return True, _info_from_meta(sig_meta)


# ── Firma encadenada (ruta de firmas) ────────────────────────────────────────

_MSIGNED_EXT = ".msigned"


def sign_for_route(username, file_path, private_key_path=None):
    """
    Firma un documento dentro de una ruta multi-firma.

    Para DOCX: añade un bloque de firma adicional al documento, encadenado
    criptográficamente con las firmas anteriores.
    Para otros formatos: usa el contenedor .msigned.

    Devuelve (output_path, output_bytes).
    """
    if private_key_path:
        with open(private_key_path, "rb") as f:
            private_key = serialization.load_pem_private_key(f.read(), password=None)
    else:
        private_key = _load_private_key(username)
    certificate = _load_certificate(username)
    ext = os.path.splitext(file_path)[1].lower()

    if ext in (".docx", ".doc"):
        return _sign_docx_route_step(username, file_path, private_key, certificate)
    else:
        return _sign_msigned_route_step(username, file_path, private_key, certificate)


def _sign_docx_route_step(username, file_path, private_key, certificate):
    """Añade un bloque de firma de ruta a un DOCX, encadenado con los anteriores."""
    from docx import Document

    # Hash del contenido original (se excluyen bloques de firma existentes)
    original_hash = _docx_content_hash(file_path)

    # Extraer firmas de ruta previas embebidas en el documento
    doc_read  = Document(file_path)
    prev_sigs = []
    for para in doc_read.paragraphs:
        if para.text.startswith(_DOCX_JSON_MARK):
            try:
                meta = json.loads(
                    base64.b64decode(para.text[len(_DOCX_JSON_MARK):]).decode("utf-8")
                )
                prev_sigs.append(meta)
            except Exception:
                pass
    prev_sigs.sort(key=lambda s: s.get("order", 1))

    # Hash de cadena
    if not prev_sigs:
        hash_chain = original_hash          # primer firmante: igual al hash de contenido
    else:
        chain_input = original_hash.encode()
        for s in prev_sigs:
            chain_input += base64.b64decode(s["signature"])
        hash_chain = hashlib.sha256(chain_input).hexdigest()

    raw_sig = _rsa_sign(private_key, hash_chain.encode())

    # Nombre original sin sufijos _firmado acumulados
    orig_name = re.sub(r"(_firmado)+$", "", os.path.splitext(os.path.basename(file_path))[0])

    sig_meta = {
        "route_signature":   True,
        "order":             len(prev_sigs) + 1,
        "signer":            username,
        "original_filename": orig_name + ".docx",
        "algorithm":         "RSA-PSS-SHA256",
        "hash_sha256":       original_hash,
        "hash_chain":        hash_chain,
        "signature":         base64.b64encode(raw_sig).decode(),
        "certificate":       certificate,
        "signed_at":         datetime.now(_CDMX).strftime("%Y-%m-%dT%H:%M:%S"),
    }

    # Ruta de salida: siempre {original}_firmado.docx
    base_stem   = re.sub(r"(_firmado)+$", "", os.path.splitext(file_path)[0])
    signed_path = base_stem + "_firmado.docx"

    # Escribir en temporal para evitar leer y escribir el mismo archivo
    tmp_out = signed_path + ".writing"
    shutil.copy2(file_path, tmp_out)
    doc = Document(tmp_out)
    _append_docx_sig_block(doc, sig_meta, certificate)
    doc.save(tmp_out)
    os.replace(tmp_out, signed_path)

    with open(signed_path, "rb") as f:
        output_bytes = f.read()
    return signed_path, output_bytes


def _sign_msigned_route_step(username, file_path, private_key, certificate):
    """Firma de ruta para formatos no-DOCX usando contenedor .msigned."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == _MSIGNED_EXT:
        with zipfile.ZipFile(file_path, "r") as zf:
            original_bytes = zf.read("original")
            prev_data      = json.loads(zf.read("signatures.json").decode("utf-8"))
    else:
        with open(file_path, "rb") as f:
            original_bytes = f.read()
        prev_data = {
            "version":           2,
            "original_filename": os.path.basename(file_path),
            "route":             [],
            "signatures":        [],
        }

    prev_sigs = prev_data.get("signatures", [])

    chain_input = original_bytes
    for s in prev_sigs:
        chain_input += base64.b64decode(s["signature"])
    hash_chain    = hashlib.sha256(chain_input).hexdigest()
    hash_original = hashlib.sha256(original_bytes).hexdigest()

    raw_sig = _rsa_sign(private_key, hash_chain.encode())
    new_sig = {
        "order":       len(prev_sigs) + 1,
        "signer":      username,
        "signed_at":   datetime.now(_CDMX).strftime("%Y-%m-%dT%H:%M:%S"),
        "hash_sha256": hash_original,
        "hash_chain":  hash_chain,
        "signature":   base64.b64encode(raw_sig).decode(),
        "certificate": certificate,
    }

    output_data = {
        "version":           2,
        "original_filename": prev_data.get("original_filename", os.path.basename(file_path)),
        "route":             prev_data.get("route", []),
        "signatures":        prev_sigs + [new_sig],
    }

    base_stem   = re.sub(r"(_firmado)+$", "", os.path.splitext(file_path)[0])
    output_path = base_stem + "_firmado.msigned"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("original",
                    original_bytes)
        zf.writestr("signatures.json",
                    json.dumps(output_data, indent=2, ensure_ascii=False).encode("utf-8"))
    output_bytes = buf.getvalue()
    with open(output_path, "wb") as f:
        f.write(output_bytes)
    return output_path, output_bytes


def verify_route_file(file_path):
    """
    Verifica un contenedor .msigned de ruta multi-firma.
    Comprueba:
      1. El orden de firmas coincide con la ruta embebida.
      2. Cada hash de cadena es correcto (el orden no fue alterado).
      3. Cada firma RSA es válida criptográficamente.
    """
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            original_bytes = zf.read("original")
            data = json.loads(zf.read("signatures.json").decode("utf-8"))
    except Exception:
        return False, "El archivo no es un contenedor multi-firma válido (.msigned)."

    signatures = data.get("signatures", [])
    route      = data.get("route",      [])

    if not signatures:
        return False, "El documento no contiene ninguna firma."

    # Verificar que el orden de firmantes coincide con la ruta definida
    if route:
        signers = [s["signer"] for s in sorted(signatures, key=lambda x: x["order"])]
        if signers != route:
            return False, (
                "Las firmas NO siguen la ruta definida.\n\n"
                f"Ruta esperada  : {' → '.join(route)}\n"
                f"Firmas halladas: {' → '.join(signers)}"
            )

    # Verificar cada firma en la cadena
    chain_input   = original_bytes
    hash_original = hashlib.sha256(original_bytes).hexdigest()

    for i, sig in enumerate(sorted(signatures, key=lambda x: x["order"])):
        expected_hash = hashlib.sha256(chain_input).hexdigest()
        if sig.get("hash_chain", "") != expected_hash:
            return False, (
                f"La firma #{sig['order']} ({sig.get('signer','?')}) no es válida.\n"
                "El hash de cadena no coincide: el orden fue alterado o falta una firma."
            )
        try:
            pub_key = _public_key_from_cert(sig.get("certificate", {}))
            raw_sig = base64.b64decode(sig["signature"])
            _rsa_verify(pub_key, raw_sig, expected_hash.encode())
        except ValueError as e:
            return False, str(e)
        except Exception:
            return False, (
                f"La firma #{sig['order']} ({sig.get('signer','?')}) "
                "no es válida criptográficamente."
            )
        chain_input += raw_sig

    last = sorted(signatures, key=lambda x: x["order"])[-1]
    all_signers = " → ".join(
        f"{s['signer']} (#{s['order']})"
        for s in sorted(signatures, key=lambda x: x["order"])
    )
    return True, {
        "signer":            last["signer"],
        "original_filename": data.get("original_filename", ""),
        "signed_at":         last["signed_at"],
        "hash_sha256":       hash_original,
        "cert_status":       last.get("certificate", {}).get("status", "unknown"),
        "route":             route,
        "all_signers":       all_signers,
        "total_signatures":  len(signatures),
    }


# ── API pública ───────────────────────────────────────────────────────────────

def sign_file(username, file_path, private_key_path=None):
    if private_key_path:
        with open(private_key_path, "rb") as f:
            private_key = serialization.load_pem_private_key(f.read(), password=None)
    else:
        private_key = _load_private_key(username)
    certificate = _load_certificate(username)
    ext = os.path.splitext(file_path)[1].lower()
    if ext in (".docx", ".doc"):
        return _sign_docx(username, file_path, private_key, certificate)
    elif ext == ".pdf":
        return _sign_pdf(username, file_path, private_key, certificate)
    else:
        return _sign_container(username, file_path, private_key, certificate)


def verify_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == _MSIGNED_EXT:
        return verify_route_file(file_path)
    elif ext in (".docx", ".doc"):
        return _verify_docx(file_path)
    elif ext == ".pdf":
        return _verify_pdf(file_path)
    else:
        return _verify_container(file_path)


# ── Firma de mensajes de texto ────────────────────────────────────────────────

def sign_message(username, message):
    dirs = get_user_dir(username)
    private_path = os.path.join(dirs["keys"], f"{username}_private.pem")
    with open(private_path, "rb") as f:
        private_key = serialization.load_pem_private_key(f.read(), password=None)

    signature = private_key.sign(
        message.encode(),
        padding.PSS(
            mgf=padding.MGF1(hashes.SHA256()),
            salt_length=padding.PSS.MAX_LENGTH,
        ),
        hashes.SHA256(),
    )

    timestamp      = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    signature_path = os.path.join(dirs["sign"], f"signature_{timestamp}.sig")
    with open(signature_path, "wb") as f:
        f.write(signature)
    return signature_path


def verify_signature(username, message, signature_file):
    dirs = get_user_dir(username)
    public_path = os.path.join(dirs["keys"], f"{username}_public.pem")
    with open(public_path, "rb") as f:
        public_key = serialization.load_pem_public_key(f.read())

    with open(signature_file, "rb") as f:
        signature = f.read()

    try:
        public_key.verify(
            signature,
            message.encode(),
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),
                salt_length=padding.PSS.MAX_LENGTH,
            ),
            hashes.SHA256(),
        )
        return True
    except Exception:
        return False
